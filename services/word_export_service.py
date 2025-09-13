import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from io import BytesIO
from config import EXPORT_CONFIG, DEFAULT_COLORS
from services.color_service import get_course_color
from services.base_export_service import BaseExportService

class WordExportService(BaseExportService):
    """Word export service"""
    
    def generate_word(self, df, output, user_selected_colors, title):
        """Generate Word file"""
        try:
            # Create document
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = 1  # Center alignment
            
            # Add table (8 columns: period/weekday + 7 days)
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            
            # Set headers
            headers = ['节次/星期', '星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                # Set header style
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.alignment = 1  # Center alignment
                    run = paragraph.runs[0]
                    run.font.bold = True
            
            # Create course dictionary for quick lookup
            course_dict = {}
            for _, course in df.iterrows():
                day = course.get('星期', '')
                period = course.get('节次', '')
                if day and period:
                    key = (day, period)
                    if key not in course_dict:
                        course_dict[key] = []
                    course_dict[key].append(course)
            
            # Keep track of assigned colors to ensure different courses have different colors
            assigned_colors = {}
            
            # Fill course data
            # Add data for morning, afternoon, and evening study separately
            periods_info = [
                {'name': '上午', 'start': 1, 'count': 4},
                {'name': '下午', 'start': 5, 'count': 4},
                {'name': '晚自习', 'start': 9, 'count': 4}
            ]
            
            row_index = 1  # Start from the second row (first row is header)
            for period_info in periods_info:
                # Add period identifier row
                row_cells = table.add_row().cells
                period_cell = row_cells[0]
                period_cell.merge(row_cells[7])  # Merge all columns
                period_cell.text = period_info['name']
                
                # Set period identifier style
                for paragraph in period_cell.paragraphs:
                    paragraph.alignment = 1  # Center alignment
                    run = paragraph.runs[0]
                    run.font.bold = True
                
                # Add courses for this period
                for i in range(period_info['count']):
                    row_cells = table.add_row().cells
                    period_num = period_info['start'] + i
                    
                    # Set period
                    row_cells[0].text = f'第{period_num}节'
                    for paragraph in row_cells[0].paragraphs:
                        paragraph.alignment = 1  # Center alignment
                    
                    # Fill daily courses
                    days = ['星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']
                    for col, day in enumerate(days, 1):
                        # Convert "星期一" etc. to "周一" etc. to match data
                        day_short = day.replace('星期', '周')
                        key = (day_short, period_num)
                        
                        if key in course_dict:
                            courses = course_dict[key]
                            # If there are multiple courses at the same time, display all courses
                            course_texts = []
                            for course in courses:
                                course_name = course.get('课程名称', '')
                                teacher = course.get('教师', '')
                                location = course.get('地点', '')
                                start_time = course.get('开始时间', '')
                                end_time = course.get('结束时间', '')
                                notes = course.get('备注', '')
                                
                                course_text = course_name
                                if teacher and teacher != '未指定':
                                    course_text += f"\n教师：{teacher}"
                                if location and location != '未指定':
                                    course_text += f"\n地点：{location}"
                                if start_time and end_time:
                                    course_text += f"\n时间：{start_time}~{end_time}"
                                if notes:
                                    course_text += f"\n备注：{notes}"
                                course_texts.append(course_text)
                            
                            row_cells[col].text = '\n'.join(course_texts)
                            # Set course cell style
                            for paragraph in row_cells[col].paragraphs:
                                paragraph.alignment = 1  # Center alignment
                            
                            # Apply background color
                            first_course = courses[0]
                            course_name = first_course.get('课程名称', '')
                            if course_name:
                                # Check if we've already assigned a color for this course name
                                if course_name in assigned_colors:
                                    # Use the already assigned color for consistency
                                    color_rgb = assigned_colors[course_name]
                                else:
                                    # Get color from color service
                                    color_rgb = get_course_color(course_name, user_selected_colors)
                                    
                                    # Ensure this color is not already used by another course
                                    used_colors = set(assigned_colors.values())
                                    color_index = 0
                                    # If the color is already used, find a different one
                                    while tuple(color_rgb) in used_colors and color_index < len(DEFAULT_COLORS):
                                        color_rgb = DEFAULT_COLORS[color_index]
                                        color_index += 1
                                    
                                    # Store the assigned color for this course name
                                    assigned_colors[course_name] = color_rgb
                                
                                # Apply background color in Word
                                shading_el = parse_xml(
                                    f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}"/>'
                                )
                                row_cells[col]._tc.get_or_add_tcPr().append(shading_el)
                    
                row_index += 1
            
            # Set column widths
            for i, column in enumerate(table.columns):
                for cell in column.cells:
                    cell.width = Inches(1.2)
            
            # Save to output stream
            doc.save(output)
        except Exception as e:
            # Record detailed error information
            import traceback
            error_info = traceback.format_exc()
            print(f"Error generating Word: {str(e)}")
            print(f"Detailed error information:\n{error_info}")
            raise e

    def create_word_export(self, data):
        """Create Word export"""
        df, title, raw_data = self.prepare_data(data)
        user_selected_colors = raw_data.get('userSelectedColors', {})
        
        # Generate Word file to memory
        output = BytesIO()
        self.generate_word(df, output, user_selected_colors, title)
        output.seek(0)
        
        # Return file stream directly
        filename = f"{title}.docx"
        return output, filename