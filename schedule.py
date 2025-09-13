import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
import json
import random
from flask import Flask, request, jsonify, send_file, render_template
import os
# 添加必要的导入
from io import BytesIO

app = Flask(__name__, static_folder='static', template_folder='templates')

# 颜色映射规则（基于需求分析报告）
COLOR_MAP = {
    "语文": (255, 204, 204),    # 粉色
    "数学": (204, 255, 255),    # 浅蓝色
    "英语": (204, 255, 204),    # 绿色
    "综研": (229, 229, 204),    # 浅绿色
    "趣味体育": (255, 255, 153), # 黄色
    "体育": (153, 204, 255),    # 蓝色
    "音乐": (221, 170, 221),    # 浅紫色
    "体育与健康": (153, 204, 255), # 浅蓝色
    "道法": (204, 153, 204),    # 紫色
    "美术": (255, 179, 136),    # 橙色
    "科学": (255, 229, 153),    # 浅黄色
}

# 额外的默认颜色，用于自动分配给未定义颜色的课程
DEFAULT_COLORS = [
    (255, 192, 203),  # 粉色
    (173, 216, 230),  # 浅蓝
    (144, 238, 144),  # 浅绿
    (255, 182, 193),  # 浅粉红
    (221, 160, 221),  # 梅花色
    (175, 238, 238),  # 浅青色
    (255, 218, 185),  # 桃色
    (240, 230, 140),  # 卡其色
    (230, 230, 250),  # 薰衣草色
    (255, 228, 196),  # 比卡迪色
    (255, 160, 122),  # 浅珊瑚色
    (176, 224, 230),  # 粉蓝
    (255, 228, 181),  # 麦色
    (189, 183, 107),  # 暗卡其色
    (216, 191, 216),  # 苍紫罗兰色
    (152, 251, 152),  # 薄荷奶油色
    (173, 216, 230),  # 天蓝色
    (255, 192, 203),  # 粉红色
    (244, 164, 96),   # 沙棕色
    (210, 180, 140),  # 萨摩色
    (255, 215, 0),    # 金色
    (218, 112, 214),  # 兰花色
    (192, 192, 192),  # 灰色
    (128, 128, 0),    # 橄榄色
    (128, 0, 128),    # 紫色
    (0, 128, 128),    # 水鸭色
    (0, 0, 128),      # 海军蓝
    (139, 0, 0),      # 深红色
    (0, 100, 0),      # 深绿色
    (128, 0, 0),      # 栗色
]

def get_course_color(course_name, user_selected_colors=None):
    """
    获取课程颜色，优先级：
    1. 用户选择的颜色
    2. 预定义的颜色映射
    3. 自动分配默认颜色（基于课程名称严格计算）
    """
    # 如果用户选择了颜色，则优先使用用户选择的颜色
    if user_selected_colors and course_name in user_selected_colors:
        color = user_selected_colors[course_name]
        # 如果颜色是字符串格式，转换为元组
        if isinstance(color, str):
            return tuple(map(int, color.split(',')))
        return color
    
    # 如果预定义了颜色，则使用预定义颜色
    if course_name in COLOR_MAP:
        return COLOR_MAP[course_name]
    
    # 基于课程名称生成稳定的哈希值，确保相同名称的课程总是获得相同的颜色
    # 使用更大的颜色池以减少不同名称课程获得相同颜色的概率
    hash_value = hash(course_name) % (len(DEFAULT_COLORS) * 100)  # 扩大颜色空间
    color_index = hash_value % len(DEFAULT_COLORS)
    return DEFAULT_COLORS[color_index]

# 用于存储课程数据的全局变量
courses_data_store = []

# 根路由，返回前端页面
@app.route('/')
def index():
    return render_template('schedule.html')

# API路由

@app.route('/api/courses', methods=['GET'])
def get_courses():
    """获取所有课程"""
    return jsonify(courses_data_store)

@app.route('/api/courses', methods=['POST'])
def add_course():
    """添加课程"""
    try:
        course_data = request.json
        courses_data_store.append(course_data)
        return jsonify({"success": True, "message": "课程添加成功"})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 400


@app.route('/api/courses/conflicts', methods=['POST'])
def check_conflicts():
    """检查课程冲突"""
    try:
        from io import StringIO
        import sys
        
        # 保存原始的stdout
        old_stdout = sys.stdout
        sys.stdout = StringIO()
        
        data = request.json
        courses = data.get('courses', [])
        df = pd.DataFrame(courses)
        conflicts = detect_conflicts(df)
        
        # 恢复原始的stdout
        sys.stdout = old_stdout
        
        return jsonify({"conflicts": conflicts})
    except Exception as e:
        return jsonify({"conflicts": [], "error": str(e)}), 400

@app.route('/api/export/excel', methods=['POST'])
def export_excel():
    """导出为Excel"""
    try:
        data = request.json
        if not data:
            return jsonify({"success": False, "message": "请求数据为空"}), 400
            
        courses = data.get('courses', [])
        title = data.get('title', '课程表')  # 从请求中获取标题
        
        # 检查课程数据
        if not courses:
            return jsonify({"success": False, "message": "没有课程数据可供导出"}), 400
        
        # 转换为DataFrame
        df = pd.DataFrame(courses)
        
        # 生成Excel文件到内存
        output = BytesIO()
        generate_excel(df, output, title)
        output.seek(0)
        
        # 直接返回文件流
        filename = f"{title}.xlsx"
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    except Exception as e:
        # 记录详细的错误信息
        import traceback
        error_info = traceback.format_exc()
        print(f"导出Excel时发生错误: {str(e)}")
        print(f"详细错误信息:\n{error_info}")
        return jsonify({"success": False, "message": f"导出失败: {str(e)}"}), 500

@app.route('/api/export/word', methods=['POST'])
def export_word():
    """导出为Word"""
    try:
        data = request.json
        courses = data.get('courses', [])
        user_selected_colors = data.get('userSelectedColors', {})
        title = data.get('title', '课程表')  # 从请求中获取标题
        df = pd.DataFrame(courses)
        
        # 生成Word文件到内存
        output = BytesIO()
        generate_word(df, output, user_selected_colors, title)
        output.seek(0)
        
        # 直接返回文件流
        filename = f"{title}.docx"
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 400

def detect_conflicts(course_data):
    """检测课程冲突：同一教师在同一时间的课程安排"""
    conflicts = []
    # 按教师和时间分组检查
    if "教师" in course_data.columns and "星期" in course_data.columns and "节次" in course_data.columns:
        grouped = course_data.groupby(["教师", "星期", "节次"])
        for key, group in grouped:
            if len(group) > 1:
                conflicts.append(f"冲突：教师{key[0]}在{key[1]}{key[2]}节有{len(group)}门课程")
    
    # 检查同一班级在同一时间的课程安排
    if "班级" in course_data.columns and "星期" in course_data.columns and "节次" in course_data.columns:
        grouped = course_data.groupby(["班级", "星期", "节次"])
        for key, group in grouped:
            if len(group) > 1:
                conflicts.append(f"冲突：班级{key[0]}在{key[1]}{key[2]}节有{len(group)}门课程")
    
    return conflicts

def generate_excel(df, output, title):
    """生成Excel文件"""
    try:
        # 创建工作簿和工作表
        wb = Workbook()
        ws = wb.active
        ws.title = title
        
        # 创建一个字典来快速查找课程
        course_dict = {}
        for _, course in df.iterrows():
            # 确保星期和节次字段存在
            day = course.get('星期', '')
            period = course.get('节次', '')
            if day and period:
                key = (day, period)
                if key not in course_dict:
                    course_dict[key] = []
                course_dict[key].append(course)
        
        # 设置标题
        ws.merge_cells('A1:H1')
        title_cell = ws['A1']  # 获取合并区域的左上角单元格
        title_cell.value = title
        title_cell.font = Font(size=16, bold=True)
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # 设置表头
        headers = ['节次/星期', '星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=2, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        
        # 设置上午/下午/晚自习标识和课程
        time_periods = ['上午', '下午', '晚自习']
        # 每个时段开始行: 上午从第3行开始，下午从第8行开始，晚自习从第13行开始
        period_starts = [3, 8, 13]
        
        for i, (period, start_row) in enumerate(zip(time_periods, period_starts)):
            # 设置时段标识（合并单元格）
            ws.merge_cells(start_row=start_row, start_column=1, end_row=start_row, end_column=8)
            period_cell = ws[f'A{start_row}']  # 获取合并区域的左上角单元格
            period_cell.value = period
            period_cell.font = Font(bold=True)
            period_cell.alignment = Alignment(horizontal='center', vertical='center')
            period_cell.fill = PatternFill(start_color="E2E8F0", end_color="E2E8F0", fill_type="solid")
            
            # 填充该时段的课程（4节课）
            for row_offset in range(4):
                row = start_row + row_offset + 1  # 时段标识行+1开始填课程
                period_num = row_offset + 1 + i * 4  # 节次编号：1-4, 5-8, 9-12
                # 设置节次
                ws.cell(row=row, column=1, value=f'第{period_num}节')
                
                # 填充每天的课程
                days = ['星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']
                for col, day in enumerate(days, 2):
                    # 将"星期一"等转换为"周一"等以匹配数据
                    day_short = day.replace('星期', '周')
                    key = (day_short, period_num)
                    if key in course_dict:
                        courses = course_dict[key]
                        # 如果同一时间有多门课程，显示所有课程
                        course_texts = []
                        for course in courses:
                            # 获取课程信息，确保字段存在
                            course_name = course.get('课程名称', '')
                            teacher = course.get('教师', '')
                            location = course.get('地点', '')
                            start_time = course.get('开始时间', '')
                            end_time = course.get('结束时间', '')
                            notes = course.get('备注', '')
                            
                            course_text = course_name
                            if teacher and teacher != '未指定':
                                course_text += f"\n教师：{teacher}"
                            if location and location != '未指定':
                                course_text += f"\n地点：{location}"
                            if start_time and end_time:
                                course_text += f"\n时间：{start_time}~{end_time}"
                            if notes:
                                course_text += f"\n备注：{notes}"
                            course_texts.append(course_text)
                        
                        cell = ws.cell(row=row, column=col, value='\n'.join(course_texts))
                        cell.alignment = Alignment(wrap_text=True, vertical='center')
                        
                        # 获取课程颜色
                        first_course = courses[0]
                        course_name = first_course.get('课程名称', '')
                        if course_name:
                            color_rgb = get_course_color(course_name)
                            color_hex = f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}"
                            fill = PatternFill(start_color=color_hex, end_color=color_hex, fill_type="solid")
                            cell.fill = fill
        
        # 设置列宽
        column_widths = [12, 15, 15, 15, 15, 15, 15, 15]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[chr(64 + i)].width = width
        
        # 设置行高
        for row in range(1, 18):  # 更新行范围以适应新的结构
            if row in [3, 8, 13]:  # 上午/下午/晚自习行
                ws.row_dimensions[row].height = 25
            else:
                ws.row_dimensions[row].height = 40
        
        # 保存到输出流
        wb.save(output)
    except Exception as e:
        # 记录详细的错误信息
        import traceback
        error_info = traceback.format_exc()
        print(f"生成Excel时发生错误: {str(e)}")
        print(f"详细错误信息:\n{error_info}")
        raise e

def generate_word(df, output, user_selected_colors, title):
    """生成Word文件"""
    try:
        # 创建文档
        doc = Document()
        
        # 添加标题
        title_para = doc.add_heading(title, 0)
        title_para.alignment = 1  # 居中对齐
        
        # 添加表格（8列：节次/星期 + 7天）
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        
        # 设置表头
        headers = ['节次/星期', '星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # 设置表头样式
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = 1  # 居中对齐
                run = paragraph.runs[0]
                run.font.bold = True
        
        # 创建课程字典以便快速查找
        course_dict = {}
        for _, course in df.iterrows():
            day = course.get('星期', '')
            period = course.get('节次', '')
            if day and period:
                key = (day, period)
                if key not in course_dict:
                    course_dict[key] = []
                course_dict[key].append(course)
        
        # 填充课程数据
        # 为上午、下午、晚自习分别添加数据
        periods_info = [
            {'name': '上午', 'start': 1, 'count': 4},
            {'name': '下午', 'start': 5, 'count': 4},
            {'name': '晚自习', 'start': 9, 'count': 4}
        ]
        
        row_index = 1  # 从第二行开始（第一行是表头）
        for period_info in periods_info:
            # 添加时段标识行
            row_cells = table.add_row().cells
            period_cell = row_cells[0]
            period_cell.merge(row_cells[7])  # 合并所有列
            period_cell.text = period_info['name']
            
            # 设置时段标识样式
            for paragraph in period_cell.paragraphs:
                paragraph.alignment = 1  # 居中对齐
                run = paragraph.runs[0]
                run.font.bold = True
            
            # 添加该时段的课程
            for i in range(period_info['count']):
                row_cells = table.add_row().cells
                period_num = period_info['start'] + i
                
                # 设置节次
                row_cells[0].text = f'第{period_num}节'
                for paragraph in row_cells[0].paragraphs:
                    paragraph.alignment = 1  # 居中对齐
                
                # 填充每天的课程
                days = ['星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']
                for col, day in enumerate(days, 1):
                    # 将"星期一"等转换为"周一"等以匹配数据
                    day_short = day.replace('星期', '周')
                    key = (day_short, period_num)
                    
                    if key in course_dict:
                        courses = course_dict[key]
                        # 如果同一时间有多门课程，显示所有课程
                        course_texts = []
                        for course in courses:
                            course_name = course.get('课程名称', '')
                            teacher = course.get('教师', '')
                            location = course.get('地点', '')
                            start_time = course.get('开始时间', '')
                            end_time = course.get('结束时间', '')
                            notes = course.get('备注', '')
                            
                            course_text = course_name
                            if teacher and teacher != '未指定':
                                course_text += f"\n教师：{teacher}"
                            if location and location != '未指定':
                                course_text += f"\n地点：{location}"
                            if start_time and end_time:
                                course_text += f"\n时间：{start_time}~{end_time}"
                            if notes:
                                course_text += f"\n备注：{notes}"
                            course_texts.append(course_text)
                        
                        row_cells[col].text = '\n'.join(course_texts)
                        # 设置课程单元格样式
                        for paragraph in row_cells[col].paragraphs:
                            paragraph.alignment = 1  # 居中对齐
                        
                        # 应用背景色
                        first_course = courses[0]
                        course_name = first_course.get('课程名称', '')
                        if course_name:
                            # 获取课程颜色
                            color_rgb = get_course_color(course_name, user_selected_colors)
                            # 在Word中应用背景色
                            shading_el = parse_xml(
                                f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}"/>'
                            )
                            row_cells[col]._tc.get_or_add_tcPr().append(shading_el)
                
                row_index += 1
        
        # 设置列宽
        for i, column in enumerate(table.columns):
            for cell in column.cells:
                cell.width = Inches(1.2)
        
        # 保存到输出流
        doc.save(output)
    except Exception as e:
        # 记录详细的错误信息
        import traceback
        error_info = traceback.format_exc()
        print(f"生成Word时发生错误: {str(e)}")
        print(f"详细错误信息:\n{error_info}")
        raise e

# 启动Flask应用
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
